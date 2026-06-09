"""Markdown → HTML and HTML → PDF rendering for resume packages.

The HTML and PDF are produced from the same source so what you see in Chrome
is exactly what lands in the PDF. Playwright (headless Chromium) handles the
PDF step so the print engine matches the browser engine.
"""

from __future__ import annotations

import html as _html_lib
import re
from pathlib import Path

# ATS-safe, recruiter-ready resume style.
# - Arial throughout: universally supported, ATS-readable.
# - All spacing via line-height and padding, never empty div spacers.
# - No position:fixed, no CSS multi-column, no CSS Grid with gap.
# - All text is real HTML text nodes (no SVG/image text).
RESUME_STYLE = """
@page { size: letter; margin: 0.7in 0.75in; }
* { box-sizing: border-box; }
body {
  font-family: Arial, sans-serif;
  font-size: 10.5pt;
  line-height: 1.45;
  color: #1a1a1a;
  max-width: 7in;
  margin: 0 auto;
  padding: 0;
}
h1 {
  font-family: Arial, sans-serif;
  font-size: 22pt;
  margin: 0 0 0.05in 0;
  font-weight: 700;
  letter-spacing: -0.5px;
}
h2 {
  font-family: Arial, sans-serif;
  font-size: 11pt;
  text-transform: uppercase;
  letter-spacing: 0.6px;
  margin: 0.22in 0 0.05in 0;
  padding-bottom: 0.03in;
  border-bottom: 1px solid #444;
  font-weight: 700;
}
h3 {
  font-family: Arial, sans-serif;
  font-size: 11pt;
  margin: 0.13in 0 0.02in 0;
  font-weight: 600;
  display: flex;
  justify-content: space-between;
  align-items: baseline;
}
h3 em {
  font-style: italic;
  font-weight: 400;
  font-size: 10pt;
  color: #555;
  white-space: nowrap;
  margin-left: 0.2in;
}
p { margin: 0.04in 0; }
p.company { margin: 0 0 0.04in 0; color: #444; font-size: 10pt; }
ul { margin: 0.05in 0 0.1in 0; padding-left: 0.22in; }
li { margin: 0.03in 0; }
a { color: #1a1a1a; text-decoration: none; }
strong { font-weight: 600; }
em { font-style: italic; color: #555; }
hr { border: none; border-top: 0.5px solid #888; margin: 0.1in 0; }
""".strip()


# ── Structured (OS-grade) resume rendering ───────────────────────────────────
# A semantic resume renderer ported from the Resume OS build pipeline. Instead
# of a generic Markdown→HTML pass, it parses the resume Markdown into a model
# (header, summary, experience roles, projects, skills, education) and emits
# classed HTML with recruiter-grade typography: right-aligned dates, ruled
# section headers, accent colors, and tight, print-tuned spacing.
#
# Robustness contract: packages are generated and shipped without human review,
# so this MUST never hard-fail or emit an empty/garbled resume. `_resume_html`
# only uses the structured path when it parses a confident resume shape (a name
# plus at least one experience role); on any doubt or exception it falls back to
# the legacy generic `markdown_to_html` renderer.
#
# Section ordering is preserved from the source document, so when `resume_tailor`
# reorders sections to surface the most relevant experience first, the rendered
# resume reflects that order. Unrecognized `##` sections render generically
# rather than being dropped.

_RESUME_CSS = """
@page { size: Letter; margin: 0.55in; }
* { box-sizing: border-box; }
html { background: #f3f4f6; }
body {
  margin: 0; color: #172033; background: #fff;
  font-family: Arial, Helvetica, sans-serif; font-size: 10.1pt; line-height: 1.32;
}
.page { max-width: 8.5in; min-height: 11in; margin: 0 auto; padding: 0.5in; background: #fff; }
a { color: inherit; text-decoration: none; }
.top { border-bottom: 1.5px solid #172033; padding-bottom: 10px; margin-bottom: 12px; }
h1 { margin: 0; font-size: 24pt; line-height: 1; font-weight: 800; color: #172033; }
.contact { margin-top: 8px; color: #4c5870; font-size: 9pt; line-height: 1.35; }
.summary { margin-top: 10px; font-size: 10.25pt; line-height: 1.32; color: #24364a; font-weight: 600; max-width: 7.3in; }
section { margin-top: 12px; break-inside: auto; }
h2 {
  margin: 0 0 7px; padding-bottom: 3px; border-bottom: 1px solid #172033; color: #172033;
  font-size: 9pt; line-height: 1.1; letter-spacing: 0.08em; text-transform: uppercase; font-weight: 800;
}
.role { margin-top: 10px; break-inside: avoid; }
.role:first-of-type { margin-top: 0; }
.role-head { margin-bottom: 4px; }
.role-main { display: flex; justify-content: space-between; gap: 16px; align-items: baseline; }
.role-left { min-width: 0; }
.role-title { font-weight: 800; color: #172033; font-size: 10.2pt; }
.company { font-weight: 800; color: #245b73; }
.meta { color: #4c5870; font-size: 8.85pt; }
.role-dates { color: #172033; font-size: 8.85pt; white-space: nowrap; }
ul { margin: 5px 0 0 0; padding-left: 14px; }
li { margin: 3px 0; padding-left: 1px; }
.project { margin-top: 8px; break-inside: avoid; }
.project:first-of-type { margin-top: 0; }
.project-name { font-weight: 800; color: #172033; }
.project-desc { margin-top: 3px; }
.project-links { margin-top: 3px; color: #4c5870; font-size: 8.85pt; }
.skills { display: grid; gap: 4px; }
.skill-line strong { color: #172033; }
.generic-block { margin: 4px 0 0 0; }
.education-item { margin-top: 8px; break-inside: avoid; }
.education-item:first-of-type { margin-top: 0; }
.education-title { font-weight: 800; color: #172033; }
.education-meta { color: #4c5870; font-size: 8.85pt; margin-top: 2px; }
.education-detail { margin-top: 2px; }
@media print {
  html, body { background: #fff; }
  .page { width: auto; min-height: auto; margin: 0; padding: 0; }
}
""".strip()

# Map source section titles (case-insensitive) to a renderer kind.
_SECTION_ALIASES = {
    "experience": "experience",
    "work experience": "experience",
    "professional experience": "experience",
    "employment": "experience",
    "profile": "summary",
    "summary": "summary",
    "professional summary": "summary",
    "projects": "projects",
    "selected projects": "projects",
    "skills": "skills",
    "core skills": "skills",
    "technical skills": "skills",
    "education": "education",
}

_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_SKILL_RE = re.compile(r"^\*\*(.+?):\*\*\s*(.+)$")   # **Label:** items
_SKILL_RE2 = re.compile(r"^\*\*(.+?)\*\*:\s*(.+)$")  # **Label**: items
_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")


def _esc(value: str) -> str:
    return _html_lib.escape(value or "", quote=True)


def _inline(value: str) -> str:
    """Escape text, then re-introduce Markdown links and bold as HTML."""
    out = _esc(value)
    out = _LINK_RE.sub(lambda m: f'<a href="{m.group(2)}">{m.group(1)}</a>', out)
    out = _BOLD_RE.sub(r"<strong>\1</strong>", out)
    return out


def _split_once(value: str, sep: str) -> tuple[str, str]:
    idx = value.find(sep)
    if idx == -1:
        return value.strip(), ""
    return value[:idx].strip(), value[idx + len(sep):].strip()


def _strip_italics(value: str) -> str:
    v = value.strip()
    if len(v) >= 2 and v.startswith("*") and v.endswith("*"):
        v = v[1:-1]
    return v.strip()


def _is_meta_line(line: str) -> bool:
    """A `*Location · Dates*` italic meta line (not a `**bold**` heading)."""
    s = line.strip()
    return s.startswith("*") and not s.startswith("**")


def _strip_tailoring_log(text: str) -> str:
    """Remove HTML comment blocks (e.g. the resume_tailor `<!-- TAILORING LOG`)."""
    return re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)


def _parse_roles(body: list[str]) -> list[dict]:
    roles: list[dict] = []
    i, n = 0, len(body)
    while i < n:
        heading = re.match(r"^###\s+(.+)$", body[i])
        if not heading:
            i += 1
            continue
        role, company = _split_once(heading[1].strip(), " — ")
        j = i + 1
        meta = ""
        if j < n and _is_meta_line(body[j]):
            meta = _strip_italics(body[j])
            j += 1
        location, dates = _split_once(meta, " · ")
        bullets: list[str] = []
        while j < n and not body[j].startswith("### "):
            s = body[j].strip()
            bullet = _BULLET_RE.match(s)
            if bullet:
                bullets.append(bullet[1].strip())
            elif s and bullets:
                # Wrapped continuation of the previous bullet: don't drop it.
                bullets[-1] = f"{bullets[-1]} {s}".strip()
            j += 1
        roles.append({
            "role": role, "company": company,
            "location": location, "dates": dates, "bullets": bullets,
        })
        i = j
    return roles


def _parse_projects(body: list[str]) -> list[dict]:
    idxs = [i for i, ln in enumerate(body) if ln.startswith("### ")]
    projects: list[dict] = []
    for k, start in enumerate(idxs):
        end = idxs[k + 1] if k + 1 < len(idxs) else len(body)
        name = body[start][4:].strip()
        content = [s.strip() for s in body[start + 1:end] if s.strip()]
        link_i = next(
            (i for i, s in enumerate(content) if re.search(r"github\.com|https?://", s)),
            -1,
        )
        pre = content if link_i == -1 else content[:link_i]
        links = "" if link_i == -1 else content[link_i]
        # Separate prose description from project bullets so bullets render as
        # a real list, not flattened into the description paragraph.
        desc_parts: list[str] = []
        bullets: list[str] = []
        for s in pre:
            bullet = _BULLET_RE.match(s)
            if bullet:
                bullets.append(bullet[1].strip())
            elif bullets:
                bullets[-1] = f"{bullets[-1]} {s}".strip()  # wrapped continuation
            else:
                desc_parts.append(s)
        projects.append({
            "name": name,
            "description": " ".join(desc_parts),
            "bullets": bullets,
            "links": links,
        })
    return projects


def _parse_skills(body: list[str]) -> list[dict]:
    skills: list[dict] = []
    for ln in body:
        # Tolerate a leading bullet marker ("- " / "* ") without eating the
        # "**" of a bold "**Label:**" prefix.
        s = re.sub(r"^[-*]\s+", "", ln.strip()).strip()
        if not s or s.startswith("#"):
            continue
        m = _SKILL_RE.match(s) or _SKILL_RE2.match(s)
        if m:
            skills.append({"label": m[1].strip(), "items": m[2].strip()})
        else:
            skills.append({"label": "", "items": s})
    return skills


def _parse_education(body: list[str]) -> list[dict]:
    idxs = [i for i, ln in enumerate(body) if ln.startswith("### ")]
    items: list[dict] = []
    for k, start in enumerate(idxs):
        end = idxs[k + 1] if k + 1 < len(idxs) else len(body)
        degree, school = _split_once(body[start][4:].strip(), " — ")
        j = start + 1
        meta = ""
        if j < end and _is_meta_line(body[j]):
            meta = _strip_italics(body[j])
            j += 1
        location, dates = _split_once(meta, " · ")
        detail = " ".join(s.strip() for s in body[j:end] if s.strip())
        items.append({
            "degree": degree, "school": school,
            "location": location, "dates": dates, "detail": detail,
        })
    return items


def _parse_resume(markdown_text: str) -> dict:
    """Parse resume Markdown into a structured model. Order of sections is kept."""
    text = _strip_tailoring_log(markdown_text)
    lines = text.split("\n")

    name = ""
    for ln in lines:
        if ln.startswith("# ") and not ln.startswith("## "):
            name = ln[2:].strip()
            break

    sec_idx = [i for i, ln in enumerate(lines) if re.match(r"^##\s+\S", ln)]
    header_end = sec_idx[0] if sec_idx else len(lines)

    contact: list[str] = []
    summary = ""
    seen_name = False
    for ln in lines[:header_end]:
        s = ln.strip()
        if not s:
            continue
        if s.startswith("# ") and not s.startswith("## "):
            seen_name = True
            continue
        if not seen_name or s in ("---", "***", "___"):
            continue
        bold = re.match(r"^\*\*(.+?)\*\*$", s)
        if bold and not summary:
            summary = bold[1].strip()
            continue
        contact.append(s)

    sections: list[dict] = []
    for k, start in enumerate(sec_idx):
        end = sec_idx[k + 1] if k + 1 < len(sec_idx) else len(lines)
        title = re.match(r"^##\s+(.+)$", lines[start])[1].strip()
        body = lines[start + 1:end]
        kind = _SECTION_ALIASES.get(title.lower())
        if kind == "summary":
            txt = " ".join(s.strip() for s in body if s.strip())
            if txt:
                summary = f"{summary} {txt}".strip() if summary else txt
            continue
        if kind == "experience":
            sections.append({"kind": "experience", "title": title, "roles": _parse_roles(body)})
        elif kind == "projects":
            sections.append({"kind": "projects", "title": title,
                             "projects": _parse_projects(body), "lines": body})
        elif kind == "skills":
            sections.append({"kind": "skills", "title": title, "skills": _parse_skills(body)})
        elif kind == "education":
            sections.append({"kind": "education", "title": title,
                             "items": _parse_education(body), "lines": body})
        else:
            sections.append({"kind": "generic", "title": title, "lines": body})

    return {"name": name, "contact": contact, "summary": summary, "sections": sections}


def _render_generic(lines: list[str]) -> str:
    out: list[str] = []
    para: list[str] = []
    bullets: list[str] = []

    def flush_para() -> None:
        if para:
            out.append(f'<p class="generic-block">{_inline(" ".join(para))}</p>')
            para.clear()

    def flush_bullets() -> None:
        if bullets:
            out.append("<ul>" + "".join(f"<li>{_inline(b)}</li>" for b in bullets) + "</ul>")
            bullets.clear()

    for ln in lines:
        s = ln.strip()
        if not s or s in ("---", "***", "___"):
            flush_para()
            flush_bullets()
            continue
        bullet = _BULLET_RE.match(s)
        if bullet:
            flush_para()
            bullets.append(bullet[1].strip())
        else:
            flush_bullets()
            para.append(s)
    flush_para()
    flush_bullets()
    return "".join(out)


def _render_role(r: dict) -> str:
    left = f'<span class="role-title">{_esc(r["role"])}</span>'
    if r["company"]:
        left += f' <span class="company">{_esc(r["company"])}</span>'
    if r["location"]:
        left += f'<span class="meta"> · {_esc(r["location"])}</span>'
    dates = f'<div class="role-dates">{_esc(r["dates"])}</div>' if r["dates"] else ""
    head = (f'<div class="role-head"><div class="role-main">'
            f'<div class="role-left">{left}</div>{dates}</div></div>')
    bullets = ""
    if r["bullets"]:
        bullets = "<ul>" + "".join(f"<li>{_inline(b)}</li>" for b in r["bullets"]) + "</ul>"
    return f'<article class="role">{head}{bullets}</article>'


def _render_section(sec: dict) -> str:
    h2 = f'<h2>{_esc(sec["title"])}</h2>'
    kind = sec["kind"]
    if kind == "experience":
        roles = "".join(_render_role(r) for r in sec["roles"])
        return f"<section>{h2}{roles}</section>"
    if kind == "projects":
        if not sec["projects"]:
            return f'<section>{h2}{_render_generic(sec["lines"])}</section>'
        blocks = []
        for p in sec["projects"]:
            block = f'<div class="project"><div class="project-name">{_esc(p["name"])}</div>'
            if p["description"]:
                block += f'<div class="project-desc">{_inline(p["description"])}</div>'
            if p.get("bullets"):
                block += "<ul>" + "".join(f"<li>{_inline(b)}</li>" for b in p["bullets"]) + "</ul>"
            if p["links"]:
                block += f'<div class="project-links">{_inline(p["links"])}</div>'
            blocks.append(block + "</div>")
        return f'<section>{h2}{"".join(blocks)}</section>'
    if kind == "skills":
        rows = []
        for s in sec["skills"]:
            if s["label"]:
                rows.append(f'<div class="skill-line"><strong>{_esc(s["label"])}:</strong> {_inline(s["items"])}</div>')
            else:
                rows.append(f'<div class="skill-line">{_inline(s["items"])}</div>')
        return f'<section>{h2}<div class="skills">{"".join(rows)}</div></section>'
    if kind == "education":
        if not sec["items"]:
            # Education written as bold/plain lines rather than ### headings:
            # render generically rather than silently dropping the section.
            return f'<section>{h2}{_render_generic(sec["lines"])}</section>'
        blocks = []
        for e in sec["items"]:
            title = _esc(e["degree"])
            if e["school"]:
                title += f' — {_esc(e["school"])}'
            block = f'<div class="education-item"><div class="education-title">{title}</div>'
            meta = " · ".join(filter(None, [e["location"], e["dates"]]))
            if meta:
                block += f'<div class="education-meta">{_esc(meta)}</div>'
            if e["detail"]:
                block += f'<div class="education-detail">{_inline(e["detail"])}</div>'
            blocks.append(block + "</div>")
        return f'<section>{h2}{"".join(blocks)}</section>'
    return f'<section>{h2}{_render_generic(sec["lines"])}</section>'


def _structured_resume_html(markdown_text: str, title: str = "Resume") -> str | None:
    """Render a resume to classed HTML, or None if it isn't a confident resume shape."""
    model = _parse_resume(markdown_text)
    has_experience = any(
        sec["kind"] == "experience" and sec["roles"] for sec in model["sections"]
    )
    if not (model["name"] and has_experience):
        return None

    header = [f'<h1>{_esc(model["name"])}</h1>']
    if model["contact"]:
        header.append('<div class="contact">'
                      + "<br>".join(_inline(c) for c in model["contact"])
                      + "</div>")
    if model["summary"]:
        header.append(f'<div class="summary">{_inline(model["summary"])}</div>')
    body = f'<header class="top">{"".join(header)}</header>'
    body += "".join(_render_section(sec) for sec in model["sections"])

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_esc(title)}</title>
<style>
{_RESUME_CSS}
</style>
</head>
<body>
<main class="page">
{body}
</main>
</body>
</html>
"""


def _resume_html(markdown_text: str, title: str = "Resume") -> str:
    """Resume-specific HTML: structured renderer with a legacy generic fallback."""
    try:
        html = _structured_resume_html(markdown_text, title=title)
        if html:
            return html
    except Exception:
        # Never let a parse edge case break package generation.
        pass
    return markdown_to_html(markdown_text, title=title)


def markdown_to_html(markdown_text: str, title: str = "Resume") -> str:
    """Render Markdown to a complete standalone HTML document.

    Uses python-markdown with the `extra` extension for tables/footnotes/etc.
    The output is a self-contained HTML file you can open in any browser.
    """
    try:
        import markdown as md
    except ImportError as e:
        raise RuntimeError(
            "python-markdown is not installed. "
            "Run `pip install markdown` inside your venv."
        ) from e

    body_html = md.markdown(
        markdown_text,
        extensions=["extra", "sane_lists"],
        output_format="html5",
    )
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
{RESUME_STYLE}
</style>
</head>
<body>
{body_html}
</body>
</html>
"""


def html_to_pdf(html_path: Path, pdf_path: Path) -> None:
    """Render an HTML file to PDF using headless Chromium via Playwright.

    The browser engine is identical to what the user sees in Chrome, so the
    PDF and the browser preview match exactly.
    """
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as e:
        raise RuntimeError(
            "playwright is not installed. Run `pip install playwright` and "
            "then `playwright install chromium` inside your venv."
        ) from e

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    file_url = html_path.resolve().as_uri()

    with sync_playwright() as p:
        browser = p.chromium.launch()
        try:
            page = browser.new_page()
            page.goto(file_url, wait_until="networkidle")
            page.pdf(
                path=str(pdf_path),
                format="Letter",
                margin={
                    "top": "0.6in",
                    "right": "0.6in",
                    "bottom": "0.6in",
                    "left": "0.6in",
                },
                print_background=True,
                prefer_css_page_size=True,
            )
        finally:
            browser.close()


def markdown_to_docx(markdown_text: str, docx_path: Path) -> None:
    """Convert Markdown text to a .docx file.

    Uses python-docx. US Letter, margins matching the HTML/PDF output.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Run `pip install python-docx`."
        ) from e

    doc = Document()

    # Set page size and margins (US Letter).
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Set default font.
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10.5)

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            for run in p.runs:
                run.font.name = "Arial"
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            for run in p.runs:
                run.font.name = "Arial"
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            for run in p.runs:
                run.font.name = "Arial"
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("<!--"):
            continue  # Skip HTML comments (tailoring log).
        elif stripped == "-->":
            continue
        else:
            doc.add_paragraph(stripped)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def render_resume(markdown_text: str, out_dir: Path, title: str = "Resume") -> dict[str, Path]:
    """Write resume.md, resume.html, and resume.pdf into `out_dir`.

    Returns a dict with paths to the three files.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    md_path = out_dir / "resume.md"
    html_path = out_dir / "resume.html"
    pdf_path = out_dir / "resume.pdf"

    md_path.write_text(markdown_text, encoding="utf-8")
    html = _resume_html(markdown_text, title=title)
    html_path.write_text(html, encoding="utf-8")
    html_to_pdf(html_path, pdf_path)

    return {"markdown": md_path, "html": html_path, "pdf": pdf_path}
